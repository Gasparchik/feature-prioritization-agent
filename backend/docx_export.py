import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

HEADER_BG = "2D3250"
HEADER_FG = "FFFFFF"
ROW_ALT   = "F0F2FF"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_font(cell, bold=False, color: str | None = None, size_pt: int = 10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size_pt)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _parse_markdown_line(doc: Document, line: str):
    stripped = line.rstrip()
    if stripped.startswith("#### "):
        doc.add_heading(stripped[5:], level=4)
    elif stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
    elif re.match(r"^[-*] ", stripped):
        p = doc.add_paragraph(stripped[2:], style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
    elif re.match(r"^\d+\. ", stripped):
        doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
    elif stripped in ("", "---"):
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        for idx, part in enumerate(re.split(r"\*\*(.+?)\*\*", stripped)):
            run = p.add_run(part)
            if idx % 2 == 1:
                run.bold = True


def generate_docx(clusters: list[dict], prd_content: str, language: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    if language == "ru":
        title_text = "Документ требований к продукту"
        backlog_title = "Приоритизированный бэклог"
        col_headers = ["#", "Функция", "Описание", "Охват", "Влияние", "Уверен.", "Трудоёмк.", "RICE"]
    else:
        title_text = "Product Requirements Document"
        backlog_title = "Prioritized Feature Backlog"
        col_headers = ["#", "Feature", "Description", "Reach", "Impact", "Confidence", "Effort", "RICE Score"]

    title_para = doc.add_heading(title_text, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading(backlog_title, level=1)

    sorted_clusters = sorted(clusters, key=lambda x: x["rice_score"], reverse=True)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], HEADER_BG)
        _set_cell_font(hdr_cells[i], bold=True, color=HEADER_FG, size_pt=9)

    for row_idx, c in enumerate(sorted_clusters):
        row_cells = table.add_row().cells
        values = [str(row_idx + 1), c["name"],
                  c["description"][:80] + ("…" if len(c["description"]) > 80 else ""),
                  str(c["rice"]["reach"]), str(c["rice"]["impact"]),
                  f"{c['rice']['confidence']}%", f"{c['rice']['effort']}", f"{c['rice_score']:.0f}"]
        for i, val in enumerate(values):
            row_cells[i].text = val
            if row_idx % 2 == 1:
                _set_cell_bg(row_cells[i], ROW_ALT)
            _set_cell_font(row_cells[i], size_pt=9)

    doc.add_paragraph()
    for line in prd_content.split("\n"):
        _parse_markdown_line(doc, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
